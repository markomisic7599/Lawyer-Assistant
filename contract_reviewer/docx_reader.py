"""Read .docx into plain text and run-level structure for mapping and LLM input."""

from __future__ import annotations

from dataclasses import dataclass, field

from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph


@dataclass
class RunInfo:
    """One text run inside a paragraph (python-docx comment/highlight anchors are run-based)."""

    text: str
    paragraph_index: int
    run_index: int
    table_path: tuple[int, int, int] | None = None  # (table_i, row_i, cell_i) or None for body


@dataclass
class ParagraphBlock:
    """A logical paragraph (body or table cell) with its runs."""

    paragraph_index: int
    text: str
    runs: list[RunInfo] = field(default_factory=list)
    table_path: tuple[int, int, int] | None = None


@dataclass
class DocumentStructure:
    """Flattened contract structure for span mapping and chunking."""

    blocks: list[ParagraphBlock]
    full_text: str

    @property
    def paragraphs(self) -> list[str]:
        return [b.text for b in self.blocks]


def _runs_from_paragraph(
    para: Paragraph,
    paragraph_index: int,
    table_path: tuple[int, int, int] | None,
) -> list[RunInfo]:
    runs: list[RunInfo] = []
    for run_index, run in enumerate(para.runs):
        runs.append(
            RunInfo(
                text=run.text or "",
                paragraph_index=paragraph_index,
                run_index=run_index,
                table_path=table_path,
            )
        )
    return runs


def _append_paragraph_block(
    blocks: list[ParagraphBlock],
    para: Paragraph,
    paragraph_index: int,
    table_path: tuple[int, int, int] | None,
) -> None:
    runs = _runs_from_paragraph(para, paragraph_index, table_path)
    text = "".join(r.text for r in runs)
    blocks.append(
        ParagraphBlock(
            paragraph_index=paragraph_index,
            text=text,
            runs=runs,
            table_path=table_path,
        )
    )


def iter_block_paragraphs(doc: Document):
    """
    Yield (paragraph, paragraph_index, table_path) in document order.
    table_path is None for body paragraphs, else (table_ord, row_i, cell_i).
    """
    paragraph_index = 0
    table_ord = 0
    for element in doc.element.body:
        tag = element.tag.split("}")[-1]
        if tag == "p":
            yield Paragraph(element, doc), paragraph_index, None
            paragraph_index += 1
        elif tag == "tbl":
            table = Table(element, doc)
            for row_i, row in enumerate(table.rows):
                for cell_i, cell in enumerate(row.cells):
                    for para in cell.paragraphs:
                        yield para, paragraph_index, (table_ord, row_i, cell_i)
                        paragraph_index += 1
            table_ord += 1


def extract_runs(doc_path: str) -> DocumentStructure:
    """
    Walk the document body and table cells in order, recording each paragraph's runs.

    paragraph_index is a monotonic index across all collected paragraphs (body + tables).
    """
    doc = Document(doc_path)
    blocks: list[ParagraphBlock] = []
    for para, paragraph_index, table_path in iter_block_paragraphs(doc):
        _append_paragraph_block(blocks, para, paragraph_index, table_path)

    full_text = "\n\n".join(b.text for b in blocks)
    return DocumentStructure(blocks=blocks, full_text=full_text)


def read_docx_to_structure(doc_path: str) -> DocumentStructure:
    """Alias for extract_runs — same as build order name."""
    return extract_runs(doc_path)


def body_paragraph_texts(doc_path: str) -> list[str]:
    """Convenience: body paragraphs only (no tables), in document order."""
    doc = Document(doc_path)
    return [p.text for p in doc.paragraphs]


def paragraph_by_flat_index(doc: Document) -> dict[int, Paragraph]:
    """Map flatten paragraph_index (body + tables) to python-docx Paragraph objects."""
    return {idx: para for para, idx, _tp in iter_block_paragraphs(doc)}
