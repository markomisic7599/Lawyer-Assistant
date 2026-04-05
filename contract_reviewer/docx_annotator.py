"""Apply highlights and reviewer notes to flagged runs in a .docx copy."""

from __future__ import annotations

from dataclasses import dataclass

from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph

from .docx_reader import paragraph_by_flat_index
from .span_mapper import RunSpan


@dataclass
class IssueAnnotation:
    """One flagged item to show in the document."""

    clause_text: str
    issue_type: str
    severity: str
    suggestion: str
    spans: list[RunSpan]


def _insert_paragraph_after(paragraph: Paragraph) -> Paragraph:
    """Create a new empty paragraph after `paragraph`, return the new paragraph."""
    new_p = OxmlElement("w:p")
    paragraph._element.addnext(new_p)
    return Paragraph(new_p, paragraph._parent)


def _highlight_run_span(paragraph: Paragraph, run_start: int, run_end: int) -> None:
    runs = paragraph.runs
    if not runs:
        return
    lo = max(0, min(run_start, len(runs) - 1))
    hi = max(0, min(run_end, len(runs) - 1))
    if lo > hi:
        lo, hi = hi, lo
    for i in range(lo, hi + 1):
        runs[i].font.highlight_color = WD_COLOR_INDEX.YELLOW


def annotate_issues(
    input_path: str,
    output_path: str,
    issues: list[IssueAnnotation],
    *,
    add_summary_page: bool = False,
) -> str:
    """
    Load input_path, apply highlights and reviewer note paragraphs, save to output_path.

    Word's native comment bubbles are not created here (python-docx has no stable
    high-level API for comments); notes appear as italic paragraphs immediately after
    the flagged paragraph. Highlights use Word's yellow highlight on affected runs.

    Returns output_path.
    """
    doc = Document(input_path)
    index_map = paragraph_by_flat_index(doc)

    # Apply highlights first (by paragraph, merge run ranges)
    for issue in issues:
        for span in issue.spans:
            para = index_map.get(span.paragraph_index)
            if para is None:
                continue
            _highlight_run_span(para, span.run_start, span.run_end)

    # One reviewer note per issue; multiple issues in the same paragraph stack notes in order
    last_note_after: dict[int, Paragraph] = {}
    for issue in issues:
        if not issue.spans:
            continue
        span = issue.spans[0]
        base = index_map.get(span.paragraph_index)
        if base is None:
            continue
        attach_after = last_note_after.get(span.paragraph_index, base)
        note = _insert_paragraph_after(attach_after)
        run = note.add_run(
            f"Reviewer note [{issue.severity}] ({issue.issue_type}): {issue.suggestion}"
        )
        run.italic = True
        if base.runs:
            run.font.size = base.runs[0].font.size
        last_note_after[span.paragraph_index] = note

    if add_summary_page:
        doc.add_paragraph("")
        h = doc.add_paragraph()
        h.add_run("Contract review summary").bold = True
        for issue in issues:
            doc.add_paragraph(
                f"[{issue.severity}] {issue.issue_type}: {issue.suggestion[:500]}"
            )

    doc.save(output_path)
    return output_path
