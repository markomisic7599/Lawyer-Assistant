"""Tests for fuzzy / exact span mapping."""

from __future__ import annotations

from pathlib import Path

from docx import Document

from contract_reviewer.docx_reader import extract_runs
from contract_reviewer.span_mapper import map_clause_to_spans, spans_for_clause_or_fallback


def test_exact_match_span(tmp_path: Path) -> None:
    p = tmp_path / "c.docx"
    doc = Document()
    doc.add_paragraph("The supplier may terminate at any time without notice.")
    doc.save(p)
    s = extract_runs(str(p))
    spans = map_clause_to_spans(s, "terminate at any time")
    assert len(spans) == 1
    assert spans[0].paragraph_index == 0


def test_fallback_paragraph(tmp_path: Path) -> None:
    p = tmp_path / "c.docx"
    doc = Document()
    doc.add_paragraph("Alpha")
    doc.add_paragraph("Beta clause here.")
    doc.save(p)
    s = extract_runs(str(p))
    spans = spans_for_clause_or_fallback(s, "no such text", fallback_paragraph_index=1)
    assert len(spans) == 1
    assert spans[0].paragraph_index == 1
