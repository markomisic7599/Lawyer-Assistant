"""Tests that annotator produces a valid .docx with extra note paragraphs."""

from __future__ import annotations

from pathlib import Path

from docx import Document

from contract_reviewer.docx_annotator import IssueAnnotation, annotate_issues
from contract_reviewer.docx_reader import extract_runs
from contract_reviewer.span_mapper import map_clause_to_spans


def test_annotate_creates_output_with_notes(tmp_path: Path) -> None:
    src = tmp_path / "in.docx"
    out = tmp_path / "out.docx"
    doc = Document()
    doc.add_paragraph("Payment is due when the moon is full.")
    doc.save(src)

    s = extract_runs(str(src))
    spans = map_clause_to_spans(s, "Payment is due")
    assert spans
    issues = [
        IssueAnnotation(
            clause_text="Payment is due",
            issue_type="ambiguity",
            severity="medium",
            suggestion="Specify a calendar date or invoice term.",
            spans=spans,
        )
    ]
    annotate_issues(str(src), str(out), issues)
    assert out.exists()
    doc2 = Document(str(out))
    texts = [p.text for p in doc2.paragraphs]
    assert any("Reviewer note" in t for t in texts)
