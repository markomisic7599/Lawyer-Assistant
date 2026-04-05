"""Tests for paragraph/run extraction."""

from __future__ import annotations

from pathlib import Path

from docx import Document

from contract_reviewer.docx_reader import extract_runs, paragraph_by_flat_index


def test_extract_runs_body(tmp_path: Path) -> None:
    p = tmp_path / "c.docx"
    doc = Document()
    doc.add_paragraph("First line.")
    doc.add_paragraph("Second line.")
    doc.save(p)

    s = extract_runs(str(p))
    assert len(s.blocks) == 2
    assert s.blocks[0].text == "First line."
    assert s.blocks[1].paragraph_index == 1
    assert len(s.blocks[0].runs) >= 1


def test_paragraph_index_map_matches_walk(tmp_path: Path) -> None:
    p = tmp_path / "c.docx"
    doc = Document()
    doc.add_paragraph("A")
    doc.save(p)
    doc2 = Document(str(p))
    m = paragraph_by_flat_index(doc2)
    assert 0 in m
    assert m[0].text == "A"
